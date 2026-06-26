#!/usr/bin/env python3
"""
inject_marks — 在 DOCX 正文指定句尾插入 [N] 手打引用标记

外部 AI 产出映射表 → 本脚本执行插入 → refcite.py --replace 收尾
本脚本不改格式、不改文字、不碰参考文献、不碰页眉页脚。

映射表格式（句级，默认）：
  P7:3:。[5]       段落7第3句句号前插[5]
  P14:6:。[1]      段落14第6句句号前插[1]
  P23:2:；[7]      段落23第2句分号前插[7]

映射表格式（段级，--para）：
  P7:[5]           段落7段尾插[5]

使用：
  # 清空所有旧 [N] 标记（重做之前调用）
  python inject_marks.py 论文.docx --strip

  # 注入新标记（从映射表）
  python inject_marks.py 论文.docx --map-file plan.txt
  python inject_marks.py 论文.docx --map "P7:3:。[5] P18:1:。[1]"
  python inject_marks.py 论文.docx --map "P7:[5]" --para

完整流程：
  python inject_marks.py 论文.docx --strip
  python inject_marks.py 论文.docx --map-file plan.txt
  python refcite.py 论文.docx --replace --verify
"""

import re, os, sys, argparse
from docx import Document
from docx.oxml.ns import qn


def parse_sentence_entry(entry):
    """解析句级条目 'P7:3:。[5]' → (7, 3, '。', 5)"""
    m = re.match(r'P(\d+):(\d+):([。！？；，])\[(\d+)\]', entry.strip())
    if not m:
        raise ValueError(f'无法解析句级条目: {entry}（格式: P7:3:。[5]）')
    return int(m.group(1)), int(m.group(2)), m.group(3), int(m.group(4))


def parse_para_entry(entry):
    """解析段级条目 'P7:[5]' → (7, 5)"""
    m = re.match(r'P(\d+):\[(\d+)\]', entry.strip())
    if not m:
        raise ValueError(f'无法解析段级条目: {entry}（格式: P7:[5]）')
    return int(m.group(1)), int(m.group(2))


def has_ref_fields(doc):
    """检测文档是否已有 REF 域（--replace 跑过）"""
    ns = qn('w:instrText')
    for p in doc.paragraphs:
        for r in p._element.findall(f'{qn("w:r")}'):
            for ins in r.findall(ns):
                if ins.text and ' REF ' in ins.text:
                    return True
    return False


def strip_marks(doc, limit=76):
    """清空正文所有 [N] 手打标记，limit 之后的段落（参考文献）不动
    
    安全：检测到 REF 域时拒绝执行（防止破坏 --replace 后的文档）。
    """
    if has_ref_fields(doc):
        print('  ⛔ 检测到 REF 域（文档已跑过 --replace），--strip 拒绝执行。')
        print('     需使用原始 DOCX（尚未跑 --replace 的版本）再操作。')
        return -1
    
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= limit: break
        for r in list(p._element.findall(f'{qn("w:r")}')):
            for t in list(r.findall(f'{qn("w:t")}')):
                old = t.text or ''
                new = re.sub(r'\[\d+\]', '', old)
                if new != old:
                    t.text = new
                    n += 1
    return n


def find_sentence_end(text, marker, sentence_num):
    """
    找 text 中第 sentence_num 个 marker 的位置（0-based）。
    返回该 marker 在 text 中的索引，找不到返段尾。
    """
    positions = [m.start() for m in re.finditer(re.escape(marker), text)]
    if len(positions) >= sentence_num:
        return positions[sentence_num - 1]
    return len(text)


def inject(doc, entry, para_mode):
    """执行一条映射条目"""
    if para_mode:
        para_idx, ref_num = parse_para_entry(entry)
        para = doc.paragraphs[para_idx]
        # 段尾追加
        if para.runs:
            para.runs[-1].text += f'[{ref_num}]'
        else:
            para.add_run(f'[{ref_num}]')
        print(f'  P{para_idx}: [{ref_num}] 段尾 ✅')
        return True
    else:
        para_idx, sentence_num, marker, ref_num = parse_sentence_entry(entry)
        para = doc.paragraphs[para_idx]
        text = para.text

        # 找第 N 个 marker 位置
        pos = find_sentence_end(text, marker, sentence_num)
        if pos >= len(text):
            # fallback 段尾
            if para.runs:
                para.runs[-1].text += f'[{ref_num}]'
            else:
                para.add_run(f'[{ref_num}]')
            print(f'  P{para_idx}: [{ref_num}] 标记「{marker}」不足{sentence_num}个，段尾 ✅')
            return True

        # 定位 run：累计偏移找到 pos 落在哪个 run
        offset = 0
        for r in para.runs:
            run_len = len(r.text)
            if offset + run_len > pos:
                local = pos - offset
                # 检查是否已有标记
                before = r.text[:local]
                if re.search(r'\[\d+\]$', before):
                    print(f'  ⏭️  P{para_idx}句{sentence_num}: 已有引用跳过')
                    return True
                r.text = before + f'[{ref_num}]' + r.text[local:]
                print(f'  P{para_idx}句{sentence_num}: [{ref_num}] ✅')
                return True
            offset += run_len

        # run 定位失败，段尾兜底
        if para.runs:
            para.runs[-1].text += f'[{ref_num}]'
        else:
            para.add_run(f'[{ref_num}]')
        print(f'  P{para_idx}: [{ref_num}] 段尾(兜底) ✅')
        return True


def main():
    parser = argparse.ArgumentParser(
        description='inject_marks — 在 DOCX 指定句尾插入 [N] 手打引用标记')
    parser.add_argument('input', help='输入 DOCX 文件路径')
    parser.add_argument('--map', help='映射表字符串（空格分隔多条）')
    parser.add_argument('--map-file', help='映射表文件（每行一条）')
    parser.add_argument('--para', action='store_true',
                        help='段级模式（默认句级）：条目格式 P7:[5]')
    parser.add_argument('--strip', action='store_true',
                        help='清空正文所有 [N] 手打标记（可单独使用或与 --map-file 搭配）')
    parser.add_argument('--output', help='输出路径（默认覆盖原文件）')
    args = parser.parse_args()

    if not args.strip and not args.map and not args.map_file:
        print('错误: 请指定 --strip、--map 或 --map-file')
        sys.exit(1)

    # 读取映射表
    entries = []
    if args.map:
        entries.extend(args.map.strip().split())
    if args.map_file:
        with open(args.map_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):
                    entries.append(line)

    mode = '段级' if args.para else '句级'
    print(f'inject_marks — 模式: {mode} | 条目: {len(entries)}')
    print(f'  输入: {args.input}')

    doc = Document(args.input)

    # --strip：清空旧标记（在注入前执行）
    if args.strip:
        n = strip_marks(doc)
        if n < 0:  # 检测到 REF 域，拒绝
            sys.exit(1)
        print(f'  清空 [N]: {n} 处')

    if not entries:
        output = args.output or args.input
        doc.save(output)
        print(f'\n✅ 输出: {output}')
        return

    ok = 0
    for entry in entries:
        try:
            if inject(doc, entry, args.para):
                ok += 1
        except ValueError as e:
            print(f'  ❌ {e}')

    output = args.output or args.input
    doc.save(output)
    print(f'\n✅ 写入 {ok}/{len(entries)} 处 | 输出: {output}')
    print(f'   下一步: python refcite.py "{output}" --replace --verify')


if __name__ == '__main__':
    main()
