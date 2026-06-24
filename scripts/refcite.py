#!/usr/bin/env python3
"""
refcite — 参考文献引用嵌入工具

将句级映射表写入 Word DOCX，生成可点击跳转的 REF 域代码。
支持三种入口：
  句级映射（模式①）：python refcite.py 论文.docx --smap "19:1:。,24:2:。"
  自动匹配（模式③）：python refcite.py 论文.docx --auto
  替换手打标记（模式D）：python refcite.py 论文.docx --replace [--keep-numbers]

依赖: python-docx (pip install python-docx)
"""

import re, os, sys, shutil, argparse
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ═══════════════════════════════════════════════════════════
# XML 工具
# ═══════════════════════════════════════════════════════════

def make_rPr_sup():
    rPr = OxmlElement('w:rPr')
    for t, v in [('vertAlign', 'superscript'), ('sz', '24'), ('szCs', '24')]:
        e = OxmlElement(f'w:{t}'); e.set(qn('w:val'), v); rPr.append(e)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '宋体'); rPr.append(rf)
    c = OxmlElement('w:color'); c.set(qn('w:val'), '000000'); rPr.append(c)
    return rPr

def make_ref_runs(bm_name, display_num):
    """生成 5-run REF 域（begin → instrText → separate → [n] → end）"""
    rs = []
    for tag, val in [
        ('begin',          None),
        ('instrText',      f' REF {bm_name} \r \\h  \\* MERGEFORMAT'),
        ('separate',       None),
        ('display',        f'[{display_num}]'),
        ('end',            None),
    ]:
        r = OxmlElement('w:r'); r.append(make_rPr_sup())
        if tag == 'begin':
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), 'begin'); r.append(fc)
        elif tag == 'instrText':
            i = OxmlElement('w:instrText')
            i.set(qn('xml:space'), 'preserve')
            i.text = val; r.append(i)
        elif tag == 'separate':
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), 'separate'); r.append(fc)
        elif tag == 'display':
            t = OxmlElement('w:t'); t.text = val; r.append(t)
        elif tag == 'end':
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), 'end'); r.append(fc)
        rs.append(r)
    return rs

def insert_after(target, runs):
    """在 target run 后依次插入 runs"""
    ip = target
    for r in runs: ip.addnext(r); ip = r

def insert_before(target, runs):
    """在 target run 前依次插入 runs"""
    for r in runs: target.addprevious(r)

def split_at_punct(run_elem, target_chars='，。'):
    """在 run 内第一个 target_chars 处拆分，返回（后半run, 'before'|'after'）"""
    texts = run_elem.findall(f'{{{NS}}}t')
    full = ''.join(t.text or '' for t in texts)
    ppos = -1
    for ch in target_chars:
        pos = full.find(ch)
        if pos != -1: ppos = pos; break
    if ppos == -1 or ppos >= len(full): return None, None
    acc = 0; tt = None
    for t in texts:
        tlen = len(t.text or '')
        if acc + tlen > ppos: tt = t; break
        acc += tlen
    if tt is None: return None, None
    old, loc = tt.text or '', ppos - acc
    tt.text = old[:loc]           # 前半不含标点
    after = old[loc:]             # 从标点开始（含标点）
    if not after:
        return run_elem, 'after'  # 标点在 run 末尾
    nr = OxmlElement('w:r')
    rp = run_elem.find(f'{{{NS}}}rPr')
    if rp is not None: nr.append(deepcopy(rp))
    nt = OxmlElement('w:t'); nt.text = after
    if tt.get(qn('xml:space')): nt.set(qn('xml:space'), 'preserve')
    nr.append(nt); run_elem.addnext(nr)
    return nr, 'before'

def find_target(para_elem, search_from=0, max_dist=30, target_chars='，。'):
    """
    从 search_from 开始找首个含 target_chars 中任一字符的 run。
    max_dist: 最大搜索距离（字符数），超出则返回 None 回落段末
    target_chars: 要搜索的标点字符集，默认 '，。'
    """
    runs = list(para_elem.findall(f'{{{NS}}}r'))
    acc = 0; si = 0; text_pos = 0
    for i, r in enumerate(runs):
        rl = sum(len(t.text or '') for t in r.findall(f'{{{NS}}}t'))
        if acc + rl >= search_from: si = i; break
        acc += rl
        text_pos += rl
    for i in range(si, len(runs)):
        if i == 0: continue
        full = ''.join(t.text or '' for t in runs[i].findall(f'{{{NS}}}t'))
        if text_pos - search_from > max_dist:
            return None, None
        text_pos += len(full)
        for ch in target_chars:
            if ch in full:
                tr, mode = split_at_punct(runs[i], target_chars)
                if tr is not None: return tr, mode
                break  # 跳出内层循环，外层继续找下一个 run
    return None, None

def has_footnote(elem):
    """检测段落是否包含真实脚注引用（排除 separator/continuation）"""
    refs = elem.findall(f'{{{NS}}}footnoteReference')
    return len(refs) > 0

# ═══════════════════════════════════════════════════════════
# 核心功能
# ═══════════════════════════════════════════════════════════

def backup_file(path):
    """备份文件（仅首次创建）"""
    base = os.path.splitext(os.path.basename(path))[0]
    backup = os.path.join(os.path.dirname(path) or '.', f'{base}.bak.docx')
    if os.path.exists(backup) and os.path.getsize(backup) == os.path.getsize(path):
        print(f'  备份跳过（已存在）: {backup}')
        return backup
    shutil.copy2(path, backup)
    assert os.path.getsize(path) == os.path.getsize(backup)
    print(f'  备份: {backup} ✅')
    return backup

def clean_refs(doc):
    """
    清理已有 REF 域。
    安全策略：由 instrText 定位 REF 域 → 反推 fldChar 删除。
    绝不使用 if fc 匹配（会误删 TOC/PAGE/HYPERLINK 等其他域代码）。
    """
    n = 0
    for p in doc.paragraphs:
        all_runs = list(p._element.findall(f'{{{NS}}}r'))
        # 第一步：找到所有包含 REF instrText 的 run
        ref_field_idxs = set()
        for idx, r in enumerate(all_runs):
            ins = r.findall(f'{{{NS}}}instrText')
            if ins and ins[0].text and ' REF ' in ins[0].text:
                # 向前找 begin fldChar
                for j in range(idx - 1, max(idx - 5, -1), -1):
                    fc = all_runs[j].find(f'{{{NS}}}fldChar')
                    if fc is not None:
                        ref_field_idxs.add(j)
                        break
                # 标记 instrText run 本身
                ref_field_idxs.add(idx)
                # 向后找 separate 和 end fldChar
                for j in range(idx + 1, min(idx + 5, len(all_runs))):
                    ref_field_idxs.add(j)
                    fc2 = all_runs[j].find(f'{{{NS}}}fldChar')
                    if fc2 is not None and fc2.get(qn('w:fldCharType')) == 'end':
                        break
        # 第二步：移除标记的 run
        for idx in sorted(ref_field_idxs, reverse=True):
            p._element.remove(all_runs[idx])
            n += 1
    return n

REF_END_MARKERS = ('附录', '附錄', 'Appendix', '致谢', '致謝',
                    'Acknowledgements', 'Acknowledgment', '鸣谢')

def get_refs(doc):
    """从文末提取参考文献列表。返回 (title_idx, ref_start_idx, ref_end_idx)。"""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() in ('【参考文献】', '参考文献', '主要参考文献', 'References', 'Bibliography'):
            title_idx = i
            # 找到第一条参考文献
            ref_start = None
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    ref_start = j
                    break
            if ref_start is None:
                return None, None, None
            # 找到参考文献结束位置（附录/致谢等标题之前）
            ref_end = len(doc.paragraphs)
            for k in range(ref_start + 1, len(doc.paragraphs)):
                t = doc.paragraphs[k].text.strip()
                if t in REF_END_MARKERS:
                    ref_end = k
                    break
                # 也检测加粗/大号字体等格式特征？先只靠标题文本匹配
            return title_idx, ref_start, ref_end
    return None, None, None

def auto_keywords(doc, refs, ref_start):
    """自动关键词匹配（模式③）"""
    # 提取每条文献的关键词
    ref_kw = []
    for r in refs:
        kw = set()
        m = re.search(r'\.\s*(.+?)\[', r)
        if m:
            title = m.group(1)
            kw.update(re.findall(r'[\u4e00-\u9fff]{2,}', title))
            kw.update(re.findall(r'[a-zA-Z]{3,}', title))
        ref_kw.append(kw)
    
    # 逐段匹配
    mapping = {}
    for i, p in enumerate(doc.paragraphs):
        if i < ref_start and len(p.text) >= 15:
            scores = []
            for ri, kw in enumerate(ref_kw):
                score = sum(p.text.count(k) for k in kw)
                if score > 0: scores.append((score, ri))
            if scores:
                scores.sort(key=lambda x: -x[0])
                if scores[0][0] >= 2:
                    mapping[i] = scores[0][1]
    return mapping

# ═══════════════════════════════════════════════════════════
# 校验
# ═══════════════════════════════════════════════════════════

def verify_output(path):
    """保存后校验：检查书签完整性和 REF 域结构"""
    try:
        doc = Document(path)
        errors = []
        bm_starts = {}  # id → name
        bm_ends = set()  # id set
        
        for p in doc.paragraphs:
            for bm in p._element.findall(f'{{{NS}}}bookmarkStart'):
                name = bm.get(qn('w:name'))
                bm_id = bm.get(qn('w:id'))
                if name and name.startswith('_Ref') and bm_id:
                    bm_starts[bm_id] = name
            for bm in p._element.findall(f'{{{NS}}}bookmarkEnd'):
                bm_id = bm.get(qn('w:id'))
                if bm_id:
                    bm_ends.add(bm_id)
        
        # 检查每个 REF 域
        ref_count = 0
        for p in doc.paragraphs:
            runs = list(p._element.findall(f'{{{NS}}}r'))
            for i, r in enumerate(runs):
                ins = r.findall(f'{{{NS}}}instrText')
                if ins and ins[0].text and ' REF ' in ins[0].text:
                    ref_count += 1
                    if i + 4 < len(runs):
                        types = []
                        for j in range(i - 1, i + 5):
                            if 0 <= j < len(runs):
                                fc = runs[j].find(f'{{{NS}}}fldChar')
                                if fc is not None:
                                    ftype = fc.get(qn('w:fldCharType'))
                                    types.append(ftype)
                                    if ftype == 'end':
                                        break  # 遇到 end 停止，避免扫到下一个域的 begin
                        expected_types = ['begin', 'separate', 'end']
                        if types != expected_types:
                            errors.append(f'  REF 域结构异常: {types}')
        
        # 检查书签对应
        for bm_id, name in bm_starts.items():
            if bm_id not in bm_ends:
                errors.append(f'  书签 {name}(id={bm_id}) 缺少 bookmarkEnd')
        
        if errors:
            return False, f'❌ {len(errors)} 个问题:\n' + '\n'.join(errors[:5])
        return True, f'✅ REF 域 {ref_count} 处，书签 {len(bm_starts)} 个，全部完整'
    except Exception as e:
        return False, f'❌ 校验异常: {e}'


def reorder_refs_and_bookmark(doc, ref_title, ref_start, ref_end, proj):
    """重排参考文献（按项目编号）并添加书签"""
    ref_items = []
    for i in range(ref_start, ref_end):
        if doc.paragraphs[i].text.strip():
            ri = len(ref_items)
            ref_items.append((doc.paragraphs[i]._element, proj[ri], ri))

    body = doc.element.body
    for e, _, _ in ref_items:
        try:
            body.remove(e)
        except:
            pass
    ref_items.sort(key=lambda x: x[1])
    ip = doc.paragraphs[ref_title]._element
    for e, pn, _ in ref_items:
        ip.addnext(e)
        ip = e

    # 更新每条参考文献的编号文本（如 [6] → [3]）
    for e, pn, ri in ref_items:
        for r in e.findall(f'{{{NS}}}r'):
            found = False
            for t in r.findall(f'{{{NS}}}t'):
                if t.text and re.match(r'^\[\d+\]', t.text):
                    t.text = re.sub(r'^\[\d+\]', f'[{pn}]', t.text)
                    found = True
                    break
            if found:
                break

    for e, pn, _ in ref_items:
        bm = OxmlElement('w:bookmarkStart')
        bm.set(qn('w:id'), str(5000 + pn))
        bm.set(qn('w:name'), f'_Ref{5000 + pn}')
        e.insert(0, bm)
        be = OxmlElement('w:bookmarkEnd')
        be.set(qn('w:id'), str(5000 + pn))
        e.append(be)


def report_unmatched(refs, ref_done, proj):
    """打印未被引用的文献列表"""
    unmatched = set(range(len(refs))) - ref_done
    if unmatched:
        print(f'\n⚠️  以下文献未被引用:')
        for ri in sorted(unmatched):
            print(f'    [{proj[ri]}] {refs[ri][:60]}...')


# ═══════════════════════════════════════════════════════════
# 命令行入口
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='refcite — 参考文献自动引用嵌入工具')
    parser.add_argument('input', help='输入 DOCX 文件路径')
    parser.add_argument('--smap', help='句级映射: "19:1:。,24:2:；,27:3:。"（外部 AI 指定精确位置，脚本直接执行）')
    parser.add_argument('--auto', action='store_true', help='自动关键词匹配（模式③，唯一使用启发式的地方）')
    parser.add_argument('--replace', action='store_true', help='模式D：正文已有 [N] 手打标记，原位替换为 REF 域（不做启发式匹配）')
    parser.add_argument('--keep-numbers', action='store_true', help='与 --replace 搭配使用：保持原文献编号，不按首次出现顺序重排')
    parser.add_argument('--verify', action='store_true', help='保存后校验：检查书签完整性和 REF 域结构')
    parser.add_argument('--no-backup', action='store_true', help='不创建备份')
    args = parser.parse_args()

    if args.keep_numbers and not args.replace:
        print('\n错误: --keep-numbers 仅在与 --replace 搭配使用时生效')
        sys.exit(1)

    if not args.smap and not args.auto and not args.replace:
        parser.print_help()
        print('\n错误: 请指定 --smap（句级映射）、--replace（替换手打标记）或 --auto（自动匹配）')
        sys.exit(1)

    input_path = args.input
    base = os.path.splitext(os.path.basename(input_path))[0]
    # 输出路径固定由输入文件名生成，保留完整原名
    output_path = os.path.join(
        os.path.dirname(input_path) or '.', f'{base}_已引用.docx')

    print(f'refcite v1.0 — 处理: {input_path}')

    # 备份
    if not args.no_backup:
        backup_file(input_path)

    # 打开文档
    doc = Document(input_path)
    n = clean_refs(doc)
    if n: print(f'  清理旧 REF: {n} 处')

    # 提取参考文献
    ref_title, ref_start, ref_end = get_refs(doc)
    if ref_title is None:
        print('错误: 未找到参考文献标题（【参考文献】/References）')
        sys.exit(1)
    refs = [doc.paragraphs[i].text.strip()
            for i in range(ref_start, ref_end)
            if doc.paragraphs[i].text.strip()]
    print(f'  参考文献: {len(refs)} 条')

    # ── --replace 模式：正文已有 [N] 手打标记，原位替换 ──
    if args.replace:
        ref_done = set()
        
        # 先扫描全段文本找 [N] 标记（跨 run 合并）
        # 构建每一段的 run_position 映射：[(run_idx, run_elem, text, char_offset), ...]
        para_matches = []  # [(段索引, 匹配起始偏移, 匹配结束偏移, 编号)]
        for i, p in enumerate(doc.paragraphs):
            if i >= ref_start: break
            full_text = p.text
            for m in re.finditer(r'\[(\d+)\]', full_text):
                ref_num = int(m.group(1))
                if 1 <= ref_num <= len(refs):
                    para_matches.append((i, m.start(), m.end(), ref_num))
        
        if not para_matches:
            print('⚠️  未找到 [N] 手打标记')
            sys.exit(1)
        
        # 编号策略：默认按首次出现顺序，--keep-numbers 保持原编号
        if args.keep_numbers:
            proj = {ri: ri + 1 for ri in range(len(refs))}
        else:
            first_seen = {}
            for pidx, _, _, ref_num in para_matches:
                ri = ref_num - 1
                if ri not in first_seen: first_seen[ri] = len(first_seen)
            proj = {ri: i+1 for i, ri in enumerate(sorted(first_seen.keys(), key=lambda r: first_seen[r]))}
            for ri in range(len(refs)):
                if ri not in proj: proj[ri] = len(proj) + 1
        
        # 逐段替换（同段内从右到左，避免前一个替换后位置偏移）
        para_matches.sort(key=lambda x: (x[0], -x[1]))
        for pidx, ms, me, ref_num in para_matches:
            ri = ref_num - 1
            pn = proj[ri]; bmn = f'_Ref{5000+pn}'
            ref_done.add(ri)
            
            # 获取该段的所有 run
            pe = doc.paragraphs[pidx]._element
            runs = list(pe.findall(f'{{{NS}}}r'))
            
            # 构建字符偏移 → run 映射
            run_map = []  # [(run_elem, t_elem, t_text, global_start, global_end)]
            global_pos = 0
            for r in runs:
                for t in r.findall(f'{{{NS}}}t'):
                    ttxt = t.text or ''
                    run_map.append((r, t, ttxt, global_pos, global_pos + len(ttxt)))
                    global_pos += len(ttxt)
            
            # 找出 [N] 跨了哪些 t 元素
            affected = [(r, te, tt, gs, ge) for (r, te, tt, gs, ge) in run_map
                        if gs < me and ge > ms]
            
            if not affected:
                continue
            
            # 处理受影响的 t 元素：移除 [N] 对应的字符
            # 第一个 t: 保留 [ms - gs, ...) 之前的部分
            first_r, first_te, first_tt, first_gs, first_ge = affected[0]
            cut_before = max(0, ms - first_gs)
            first_te.text = first_tt[:cut_before]
            
            after_text = ''
            # 最后一个 t: 保留 (... , me - gs] 之后的部分
            if len(affected) > 1:
                last_r, last_te, last_tt, last_gs, last_ge = affected[-1]
                cut_after = max(0, me - last_gs)
                last_te.text = last_tt[cut_after:]
                # 中间的 t 清空
                for j in range(1, len(affected) - 1):
                    affected[j][1].text = ''
            else:
                # 只有一个 t 元素包含全部 [N]
                cut_after = max(0, me - first_gs)
                before_text = first_tt[:cut_before]
                after_part = first_tt[cut_after:]
                first_te.text = before_text
                if after_part:
                    after_text = after_part
            
            # 找到第一个受影响 run 后的位置插 REF
            last_run_of_first = first_r
            
            # 在 affected[0] 的 run 后插入 REF 域
            # 先检查 affected[0] 所在的 run 之后是否还有内容（同一个 run 内）
            # 如果有后续文本，插在后续 run 之前；否则插在 affected 最后一个 run 之后
            
            if not after_text and len(affected) > 1:
                # 查看最后一个 t 的剩余文本
                last_te_text = affected[-1][1].text or ''
                if last_te_text:
                    after_text = last_te_text
            
            rr = make_ref_runs(bmn, pn)
            if after_text and len(affected) == 1:
                # 单 run 内有后续文本 → 创建新 run 承载 after_text，插在 REF 域之后
                # 先创建一个包含 after_text 的新 run
                after_run = OxmlElement('w:r')
                rp = first_r.find(f'{{{NS}}}rPr')
                if rp is not None:
                    after_run.append(deepcopy(rp))
                after_t = OxmlElement('w:t')
                after_t.text = after_text
                after_t.set(qn('xml:space'), 'preserve')
                after_run.append(after_t)
                # 先插 after_run，再在它前面插 REF runs
                first_r.addnext(after_run)
                for ref_run in reversed(rr):
                    first_r.addnext(ref_run)
            elif after_text:
                # 多 run 有后续文本 → 在最后一个 run 前插入 REF
                for ref_run in reversed(rr):
                    affected[-1][0].addprevious(ref_run)
            else:
                # 无后续文本 → 在最后一个 run 后插入 REF
                insert_after(affected[-1][0], rr)
            
            print(f'  P{pidx} → [{pn}] 原位替换 [N] ✅')
        
        total_refs = len(para_matches)

    if not args.replace:
        # ── smap / auto 模式 ──
        sentence_mode = False
        if args.smap:
            # ── 句级映射：外部 AI 指定精确位置，脚本直接执行，不自动定位 ──
            sentence_mode = True
            smap = {}  # {段落索引: [(文献编号0-based, 标记字符), ...]}
            all_refs_ordered = []  # 按出现顺序记录[(段落索引, 文献编号0-based)]
            for item in args.smap.split(','):
                parts = item.strip().split(':')
                pidx = int(parts[0])
                ref_id = int(parts[1]) - 1
                marker = parts[2] if len(parts) > 2 else '。'
                if pidx not in smap: smap[pidx] = []
                smap[pidx].append((ref_id, marker))
                all_refs_ordered.append((pidx, ref_id))
            # 过滤被跳过的段落
            valid_set = {pidx for pidx in range(ref_start)
                         if len(doc.paragraphs[pidx].text.strip()) >= 15}
            skipped_p = [k for k in smap if k not in valid_set]
            smap = {k: v for k, v in smap.items() if k in valid_set}
            all_refs_ordered = [(p, r) for p, r in all_refs_ordered if p in valid_set]
            if skipped_p: print(f'  ⚠️  跳过被过滤段落: {skipped_p}')
            print(f'  句级映射: {len(all_refs_ordered)} 项')
            total_refs = len(all_refs_ordered)
        else:
            # ── 自动匹配（模式③）：唯一使用启发式的地方 ──
            mapping = auto_keywords(doc, refs, ref_start)
            print(f'  自动匹配: {len(mapping)} 段')
            total_refs = len(mapping)

        # 项目编号
        if sentence_mode:
            # 句级模式：按 all_refs_ordered 的顺序分配编号
            first_app = {}
            for pidx, ri in all_refs_ordered:
                if ri not in first_app: first_app[ri] = pidx
            sorted_ri = sorted(first_app.keys(), key=lambda r: first_app[r])
            proj = {ri: i+1 for i, ri in enumerate(sorted_ri)}
        else:
            first_app = {}
            for pidx in sorted(mapping.keys()):
                ri = mapping[pidx]
                if ri not in first_app: first_app[ri] = pidx
            sorted_ri = sorted(first_app.keys(), key=lambda r: first_app[r])
            proj = {ri: i+1 for i, ri in enumerate(sorted_ri)}
        for ri in range(len(refs)):
            if ri not in proj: proj[ri] = len(proj) + 1

        # 插入 REF 域
        ref_done = set()
        if sentence_mode:
            # ── 句级模式：外部 AI 指定位置，脚本直接执行，不自动定位 ──
            for pidx, items in sorted(smap.items()):
                pe = doc.paragraphs[pidx]._element
                text = doc.paragraphs[pidx].text
                search_start = 0
                for ri, marker in items:
                    pn = proj[ri]; bmn = f'_Ref{5000+pn}'
                    rr = make_ref_runs(bmn, pn)
                    ref_done.add(ri)
                    
                    if has_footnote(pe):
                        print(f'  P{pidx} → [{pn}] 跳过（含脚注）')
                        continue
                    
                    mpos = text.find(marker, search_start)
                    if mpos >= 0:
                        # 以映射传入的 marker 为拆分依据，不固定搜索 ，。
                        tr, mode = find_target(pe, max(0, mpos - 2), target_chars=marker)
                    else:
                        tr, mode = None, None
                    
                    if tr is not None:
                        if mode == 'before': insert_before(tr, rr)
                        else: insert_after(tr, rr)
                        print(f'  P{pidx} → [{pn}] 在「{marker}」前 ✅')
                        if mpos >= 0: search_start = mpos + 1
                    else:
                        runs = list(pe.findall(f'{{{NS}}}r'))
                        insert_after(runs[-1], rr) if runs else pe.append(rr[0])
                        print(f'  P{pidx} → [{pn}] 段末')
        else:
            # ── 段落级模式：脚本自动找插入点 ──
            for pidx, ri in sorted(mapping.items()):
                pn = proj[ri]; bmn = f'_Ref{5000+pn}'
                p = doc.paragraphs[pidx]; pe = p._element
                rr = make_ref_runs(bmn, pn)
                ref_done.add(ri)

                if has_footnote(pe):
                    print(f'  P{pidx} → [{pn}] 跳过（含脚注）')
                    continue

                # 查找书名号或学者名作为触发信号
                text = p.text
                mp = None
                for m in re.finditer(r'《[^》]+》', text): mp = m.end(); break
                if mp is None:
                    for m in re.finditer(
                        r'([\u4e00-\u9fff]{2,4})(?:认为|指出|提出|记载|主张|强调)', text):
                        mp = m.end(); break

                tr, mode = None, None
                if mp is not None:
                    tr, mode = find_target(pe, mp)

                if tr is not None:
                    if mode == 'before': insert_before(tr, rr)
                    else: insert_after(tr, rr)
                    print(f'  P{pidx} → [{pn}] ✅')
                else:
                    runs = list(pe.findall(f'{{{NS}}}r'))
                    insert_after(runs[-1], rr) if runs else pe.append(rr[0])
                    print(f'  P{pidx} → [{pn}] 段末')

    # ── 共享：重排参考文献 + 加书签 ──
    reorder_refs_and_bookmark(doc, ref_title, ref_start, ref_end, proj)

    # 保存 + 校验 + 未引用报告
    doc.save(output_path)
    label = '替换' if args.replace else '引用'
    print(f'\n✅ 输出: {output_path}')
    print(f'   {label}: {total_refs} 处 | 文献: {len(refs)} 篇')
    
    if args.verify:
        v_ok, v_msg = verify_output(output_path)
        status_tag = 'PASS' if v_ok else 'FAIL'
        print(f'   校验: [{status_tag}] {v_msg}')

    report_unmatched(refs, ref_done, proj)


if __name__ == '__main__':
    main()
